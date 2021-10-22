from docx import Document

doc = Document('这是一个文档.docx')
# print(len(doc.paragraphs))
# paragraph = doc.paragraphs[1]
# runs = paragraph.runs
# print(runs)
# for run in runs:
#     print(run.text)
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        print(run.text)


# for paragraph in doc.paragraphs:
#     print(paragraph.text)
